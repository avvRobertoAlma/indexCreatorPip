# coding=utf-8
# importo gli oggetti "path" e "listdir" #
# "path" serve per gestire i percorsi #
# "listdir" serve per elencare i file contenuti in una cartella #
from os import path, listdir, mkdir
import sys

# importo la libreria mammoth
import mammoth

# importo la libreria docx-python
import docx

# importo le funzioni contenute nel file helpers.py
from helpers import filterString, Doc, Link, createLinkArray, checkProduction, convert_to_pdf

# importo la funzione hyperlink
import hyperlink

# imposto l'encoding di default a utf-8
reload(sys)
sys.setdefaultencoding('utf-8')

# verifica sviluppo o produzione
dev_mode = checkProduction()

# salvo il percorso relativo alla directory nella quale viene eseguito lo script #
if dev_mode:
        pathDir = path.dirname(path.abspath(__file__))
else:
        pathDir = path.dirname(sys.executable)

print('Benvenuti in linkGenerator versione 1.0.0\n Il programma è realizzato dall\' Avv. Roberto Alma\n, con il supporto dell\' Avv. Matteo Moscioni, dell\'Avv. Francesco Saverio Orlando e dell\' Avv. Antonio Cocco')



# inizializzo il puntatore al documento originale
atto = docx.Document(path.join(pathDir, 'atto.docx'))

# ottengo la lista dei file contenuti nella directory
listFiles = listdir(pathDir)

# ordino la lista
listFiles.sort()

# inizializzo l'array dei documenti
documents = []

# ciclo che per tutti i file della directory, verifica se l'estensione è .pdf o meno #
# se l'estensione è pdf, crea un'istanza della classe Doc e passa come parametro "url" #
#il nome del file come estensione, e come "txt" il solo nome, senza estensione #
for file in listFiles:
    if file.endswith('.pdf'):
        name = filterString(file)
        documents.append(Doc(file, name))

# informo l'utente che i documenti sono stati correttamente registrati
print('Ho analizzato i documenti pdf')

#ottengo l'array dei link
linksArray = createLinkArray(documents)

#creazione dell'indice dei documenti
print('Inizio creazione dell\'indice dei documenti')

if dev_mode:
        indice = docx.Document()
else:
        indice = docx.Document(path.join(pathDir, 'template.docx'))
        indice.add_heading('Indice atti e documenti')


# ciclo che per ogni link salvato aggiunge un elemento all'indice e aggiunge il link ipertestuale
for link in linksArray:
    p = indice.add_paragraph()
    p.style = 'List Number 2'
    try:
            new_run = p.add_run(link.name)
            hyperlink.add(p, new_run, link.url)
    except ValueError:
            print('Error: '+ link.name)
    

indice.save(path.join(pathDir,'00_indice_atti_documenti.docx'))

# creazione della lista delle occorrenze

if dev_mode:
        lista = docx.Document()
else:
        lista = docx.Document(path.join(pathDir, 'template.docx'))
        lista.add_heading('Lista delle occorrenze')

# ciclo che per ogni link salvato aggiunge un elemento all'indice e aggiunge il link ipertestuale
for link in linksArray:
    p = lista.add_paragraph()
    try:
            new_run = p.add_run('('+link.occurrence+')')
            font = new_run.font
            font.bold = True
            font.underline = True
            hyperlink.add(p, new_run, link.url)
    except ValueError:
            print('Error: '+ link.occurrence)

lista.save(path.join(pathDir,'lista_delleOccorrenze.docx'))

# informo l'utente che inizio a scansionare l'atto
print('Sto iniziando la scansione dell\'atto\n Attendere prego..')

# scansione dell'atto e aggiunta degli hyperlinks
for p in atto.paragraphs:
    for run in p.runs:
        for link in linksArray:
            found_link = link.__eq__(run.text)
            if(found_link):
                hyperlink.add(p, run, found_link)

# informo che la scansione è completata
print('Scansione completata\n Salvataggio in corso..')

atto.save(path.join(pathDir,'new-atto.docx'))

# informo che il file è stato salvato
print('Il file completo dei link ipertestuali è stato salvato come new-atto.docx\n Grazie per aver utilizzato il programma!')

# avvio la conversione del docx in html
print('Avvio conversione docx in html')
# converto il contenuto di atto.docx in html
with open(path.join(pathDir, 'atto.docx'), 'rb') as docx_file:
    result = mammoth.convert_to_html(docx_file)
    html = result.value

for index, link in enumerate(linksArray):
        if link.occurrence in html:
                html = html.replace(link.occurrence, link.href, 1)

# add stylingRules
html = html + """
        <style>
        body, p {
                font-size: 14px;
                text-align: justify;
                font: 14px/180% Arial;
        }
        h2, h3 {
                text-align: center!important;
        }
        h6 {
                text-align: right!important;
        }
        </style>
        """

html_file = open(path.join(pathDir, 'atto-convertito.html'), 'w')
html_file.write(html)
html_file.close()


convert_to_pdf(html, path.join(pathDir, 'atto-convertito.pdf'))


